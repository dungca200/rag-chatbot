import logging
from pathlib import Path
from typing import Dict

from docx import Document

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parser for Word documents using python-docx."""

    SUPPORTED_EXTENSIONS = ['.docx']

    def parse(self, file_path: str) -> Dict:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dict with content, metadata, and status
        """
        path = Path(file_path)

        if not path.exists():
            return {
                "success": False,
                "error": f"File not found: {file_path}",
                "content": "",
                "metadata": {}
            }

        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            return {
                "success": False,
                "error": f"Unsupported file type: {path.suffix}",
                "content": "",
                "metadata": {}
            }

        try:
            doc = Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            content = "\n\n".join(paragraphs)

            metadata = {
                "filename": path.name,
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "file_size": path.stat().st_size
            }

            # Extract core properties if available
            if doc.core_properties:
                if doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
                if doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author

            logger.info(f"Parsed DOCX: {path.name}, {len(doc.paragraphs)} paragraphs")

            return {
                "success": True,
                "content": content,
                "metadata": metadata,
                "error": None
            }

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "content": "",
                "metadata": {"filename": path.name}
            }


def parse_docx(file_path: str) -> Dict:
    """Convenience function for DOCX parsing."""
    parser = DOCXParser()
    return parser.parse(file_path)
